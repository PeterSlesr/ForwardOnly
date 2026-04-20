import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import os
import json
import zipfile
import shutil
import tempfile

APP_NAME = "ForwardOnly"
FWD_EXT = ".fwd"
CONTENT_FILE = "content.txt"
SETTINGS_FILE = "settings.json"

DEFAULT_SETTINGS = {
    "window_size": 5,
    "theme": "light",
    "dark_color": "green",
    "export_txt": "",
    "export_docx": ""
}

# ── Theme definitions ─────────────────────────────────────────────────────────

THEMES = {
    "light": {
        "bg":         "#d4d0c8",
        "content_bg": "#ffffff",
        "content_fg": "#000000",
        "hidden_fg":  "#cccccc",
        "menu_bg":    "#d4d0c8",
        "menu_fg":    "#000000",
        "btn_bg":     "#d4d0c8",
        "btn_fg":     "#000000",
        "status_bg":  "#d4d0c8",
        "status_fg":  "#444444",
        "relief":     tk.RAISED,
    },
    "dark_green": {
        "bg":         "#0d0d0d",
        "content_bg": "#0a0a0a",
        "content_fg": "#33ff33",
        "hidden_fg":  "#1a3d1a",
        "menu_bg":    "#1a1a1a",
        "menu_fg":    "#33ff33",
        "btn_bg":     "#1a1a1a",
        "btn_fg":     "#33ff33",
        "status_bg":  "#111111",
        "status_fg":  "#1a8c1a",
        "relief":     tk.FLAT,
    },
    "dark_amber": {
        "bg":         "#0d0a00",
        "content_bg": "#0a0800",
        "content_fg": "#ffb000",
        "hidden_fg":  "#3d2a00",
        "menu_bg":    "#1a1400",
        "menu_fg":    "#ffb000",
        "btn_bg":     "#1a1400",
        "btn_fg":     "#ffb000",
        "status_bg":  "#111000",
        "status_fg":  "#8c6200",
        "relief":     tk.FLAT,
    }
}

def get_theme(settings):
    if settings["theme"] == "light":
        return THEMES["light"]
    return THEMES[f"dark_{settings['dark_color']}"]


# ── .fwd file I/O ─────────────────────────────────────────────────────────────

def load_fwd(path):
    with zipfile.ZipFile(path, "r") as z:
        content = z.read(CONTENT_FILE).decode("utf-8")
        try:
            settings = json.loads(z.read(SETTINGS_FILE).decode("utf-8"))
        except Exception:
            settings = {}
    merged = dict(DEFAULT_SETTINGS)
    merged.update(settings)
    return content, merged

def save_fwd(path, content, settings):
    tmp = path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(CONTENT_FILE, content.encode("utf-8"))
        z.writestr(SETTINGS_FILE, json.dumps(settings, indent=2).encode("utf-8"))
    shutil.move(tmp, path)


# ── Import from TXT / DOCX ────────────────────────────────────────────────────

def import_txt(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def import_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except ImportError:
        messagebox.showerror("Missing library",
            "python-docx is not installed.\nCannot import .docx files.")
        return None
    except Exception as e:
        messagebox.showerror("Import error", f"Could not read .docx file:\n{e}")
        return None


# ── Resource path (works both dev and PyInstaller) ───────────────────────────

def resource_path(relative):
    """Get path to bundled resource."""
    import sys
    base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, relative)

def get_exe_path():
    """Return path to the running exe (or script in dev)."""
    import sys
    if getattr(sys, 'frozen', False):
        return sys.executable
    return os.path.abspath(__file__)


# ── First-run setup ───────────────────────────────────────────────────────────

SETUP_DONE_KEY = r"Software\ForwardOnly"

def is_setup_done():
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, SETUP_DONE_KEY)
        val, _ = winreg.QueryValueEx(key, "SetupDone")
        winreg.CloseKey(key)
        return val == 1
    except Exception:
        return False

def mark_setup_done():
    try:
        import winreg
        key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, SETUP_DONE_KEY)
        winreg.SetValueEx(key, "SetupDone", 0, winreg.REG_DWORD, 1)
        winreg.CloseKey(key)
    except Exception:
        pass

def first_run_setup():
    if is_setup_done():
        return
    create_desktop_shortcut()
    register_fwd_extension()
    mark_setup_done()


# ── Desktop shortcut ──────────────────────────────────────────────────────────

def create_desktop_shortcut():
    try:
        import subprocess
        exe_path = get_exe_path()
        ico_path = resource_path(os.path.join("assets", "forwardonly.ico"))
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        shortcut_path = os.path.join(desktop, f"{APP_NAME}.lnk")
        vbs = (
            'Set oWS = WScript.CreateObject("WScript.Shell")\n'
            f'sLinkFile = "{shortcut_path}"\n'
            'Set oLink = oWS.CreateShortcut(sLinkFile)\n'
            f'oLink.TargetPath = "{exe_path}"\n'
            f'oLink.IconLocation = "{ico_path}"\n'
            'oLink.Save\n'
        )
        vbs_path = os.path.join(tempfile.gettempdir(), "fo_shortcut.vbs")
        with open(vbs_path, "w") as f:
            f.write(vbs)
        subprocess.run(["wscript", vbs_path], check=False,
                       creationflags=0x08000000)
        try:
            os.remove(vbs_path)
        except Exception:
            pass
    except Exception:
        pass


# ── Register .fwd extension ───────────────────────────────────────────────────

def register_fwd_extension():
    try:
        import winreg
        exe_path = get_exe_path()
        ico_path = resource_path(os.path.join("assets", "forwardonly.ico"))

        # .fwd -> ForwardOnly.Document
        k = winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                             r"Software\Classes\.fwd")
        winreg.SetValueEx(k, "", 0, winreg.REG_SZ, "ForwardOnly.Document")
        winreg.CloseKey(k)

        # ForwardOnly.Document description
        k = winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                             r"Software\Classes\ForwardOnly.Document")
        winreg.SetValueEx(k, "", 0, winreg.REG_SZ, "ForwardOnly Project")
        winreg.CloseKey(k)

        # Icon
        k = winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                             r"Software\Classes\ForwardOnly.Document\DefaultIcon")
        winreg.SetValueEx(k, "", 0, winreg.REG_SZ, f"{ico_path},0")
        winreg.CloseKey(k)

        # Open command
        k = winreg.CreateKey(winreg.HKEY_CURRENT_USER,
                             r"Software\Classes\ForwardOnly.Document\shell\open\command")
        winreg.SetValueEx(k, "", 0, winreg.REG_SZ, f'"{exe_path}" "%1"')
        winreg.CloseKey(k)

        # Notify Windows shell to refresh icons
        try:
            import ctypes
            ctypes.windll.shell32.SHChangeNotify(0x08000000, 0, None, None)
        except Exception:
            pass

    except Exception:
        pass


# ── Export ────────────────────────────────────────────────────────────────────

def _default_export_path(fwd_path, ext, remembered_path):
    """Return remembered path if valid, else derive from fwd_path or use Untitled."""
    if remembered_path and os.path.isdir(os.path.dirname(remembered_path)):
        return remembered_path
    if fwd_path:
        base = os.path.splitext(os.path.basename(fwd_path))[0]
        folder = os.path.dirname(fwd_path)
    else:
        base = "Untitled"
        folder = os.path.expanduser("~")
    return os.path.join(folder, base + ext)

def export_txt(content, fwd_path, remembered_path):
    initial = _default_export_path(fwd_path, ".txt", remembered_path)
    path = filedialog.asksaveasfilename(
        title="Export as TXT",
        defaultextension=".txt",
        initialfile=os.path.basename(initial),
        initialdir=os.path.dirname(initial),
        filetypes=[("Text files", "*.txt")])
    if not path:
        return None
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    os.startfile(path)
    return path

def export_docx(content, fwd_path, remembered_path):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        messagebox.showerror("Missing library",
            "python-docx is not installed.")
        return None

    initial = _default_export_path(fwd_path, ".docx", remembered_path)
    path = filedialog.asksaveasfilename(
        title="Export as DOCX",
        defaultextension=".docx",
        initialfile=os.path.basename(initial),
        initialdir=os.path.dirname(initial),
        filetypes=[("Word documents", "*.docx")])
    if not path:
        return None

    doc = Document()
    for style in doc.styles:
        if style.name == "Normal":
            style.font.name = "Courier New"
            style.font.size = Pt(11)
    for para_text in content.split("\n"):
        doc.add_paragraph(para_text)
    doc.save(path)
    os.startfile(path)
    return path


# ── Main Application ──────────────────────────────────────────────────────────

class ForwardOnly:
    def __init__(self, root, open_path=None):
        self.root = root
        self.root.title(APP_NAME)
        self.root.withdraw()

        self.fwd_path = None
        self.content = ""
        self.settings = dict(DEFAULT_SETTINGS)
        self.mode = "focus"
        self.session_text = ""
        self.import_source = None

        first_run_setup()

        if open_path:
            self._open_fwd_direct(open_path)
        else:
            self._show_launcher()

    def _open_fwd_direct(self, path):
        """Open a .fwd file directly into Focus mode (e.g. from double-click)."""
        try:
            self.content, self.settings = load_fwd(path)
        except Exception as e:
            messagebox.showerror("Error", f"Could not open file:\n{e}")
            self._show_launcher()
            return
        self.fwd_path = path
        self.session_text = ""
        self.import_source = None
        self._open_main_window("focus")

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _is_saved(self):
        return self.fwd_path is not None

    def _current_content(self):
        if self.mode == "focus":
            return self.content + self.session_text
        else:
            return self.text_area.get("1.0", tk.END).rstrip("\n")

    def _has_content(self):
        return bool(self._current_content().strip())

    def _prompt_save_before_leaving(self):
        """Returns True if safe to proceed, False if user cancelled."""
        if not self._is_saved() and self._has_content():
            answer = messagebox.askyesnocancel(
                "Unsaved project",
                "This project hasn't been saved.\nSave before leaving?")
            if answer is None:
                return False
            if answer:
                self._save_as()
                if not self._is_saved():
                    return False
        return True

    # ── Launcher ──────────────────────────────────────────────────────────────

    def _show_launcher(self):
        self.launcher = tk.Toplevel(self.root)
        self.launcher.title(APP_NAME)
        self.launcher.resizable(False, False)
        self.launcher.geometry("300x200")
        self.launcher.protocol("WM_DELETE_WINDOW", self.root.destroy)

        tk.Label(self.launcher, text="ForwardOnly",
                 font=("Courier New", 16, "bold")).pack(pady=(28, 2))
        tk.Label(self.launcher, text="a forward-only writing tool",
                 font=("Courier New", 9), fg="#888").pack(pady=(0, 20))

        tk.Button(self.launcher, text="New Project", font=("Courier New", 10),
                  command=self._new_project, width=20,
                  relief=tk.RAISED).pack(pady=5)
        tk.Button(self.launcher, text="Open Project", font=("Courier New", 10),
                  command=self._open_project, width=20,
                  relief=tk.RAISED).pack(pady=5)

    def _new_project(self):
        self.fwd_path = None
        self.content = ""
        self.settings = dict(DEFAULT_SETTINGS)
        self.session_text = ""
        self.import_source = None
        self.launcher.destroy()
        self._open_main_window("focus")

    def _open_project(self):
        path = filedialog.askopenfilename(
            title="Open project",
            filetypes=[
                ("All supported files", f"*{FWD_EXT} *.txt *.docx"),
                ("ForwardOnly files", f"*{FWD_EXT}"),
                ("Text files", "*.txt"),
                ("Word documents", "*.docx"),
                ("All files", "*.*")
            ])
        if not path:
            return

        ext = os.path.splitext(path)[1].lower()

        if ext == FWD_EXT:
            try:
                self.content, self.settings = load_fwd(path)
            except Exception as e:
                messagebox.showerror("Error", f"Could not open file:\n{e}")
                return
            self.fwd_path = path
            self.session_text = ""
            self.import_source = None
            self.launcher.destroy()
            self._ask_mode()

        elif ext == ".txt":
            content = import_txt(path)
            if content is None:
                return
            self.fwd_path = None
            self.content = content
            self.settings = dict(DEFAULT_SETTINGS)
            self.session_text = ""
            self.import_source = os.path.basename(path)
            self.launcher.destroy()
            self._open_main_window("focus")

        elif ext == ".docx":
            content = import_docx(path)
            if content is None:
                return
            self.fwd_path = None
            self.content = content
            self.settings = dict(DEFAULT_SETTINGS)
            self.session_text = ""
            self.import_source = os.path.basename(path)
            self.launcher.destroy()
            self._open_main_window("focus")

        else:
            messagebox.showerror("Unsupported file",
                "Please open a .fwd, .txt, or .docx file.")

    def _ask_mode(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Open as...")
        dialog.resizable(False, False)
        dialog.geometry("260x140")
        dialog.protocol("WM_DELETE_WINDOW", self.root.destroy)

        tk.Label(dialog, text="Open in which mode?",
                 font=("Courier New", 10)).pack(pady=(24, 12))

        btn_frame = tk.Frame(dialog)
        btn_frame.pack()

        def pick(mode):
            dialog.destroy()
            self._open_main_window(mode)

        tk.Button(btn_frame, text="Review", font=("Courier New", 10),
                  command=lambda: pick("review"), width=10,
                  relief=tk.RAISED).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Focus", font=("Courier New", 10),
                  command=lambda: pick("focus"), width=10,
                  relief=tk.RAISED).pack(side=tk.LEFT, padx=8)

    # ── Main window ───────────────────────────────────────────────────────────

    def _open_main_window(self, mode):
        self.mode = mode

        self.win = tk.Toplevel(self.root)
        fname = os.path.basename(self.fwd_path) if self.fwd_path else "Untitled"
        self.win.title(f"{APP_NAME} — {fname}")
        self.win.geometry("800x520")
        self.win.protocol("WM_DELETE_WINDOW", self._on_close)

        self._build_menu()
        self._build_toolbar()
        self._build_content()
        self._build_statusbar()
        self._apply_theme()

        if mode == "focus":
            self._enter_focus()
        else:
            self._enter_review()

        # Show import notice if applicable
        if self.import_source:
            self._show_import_notice()

        self.win.focus_force()

    def _show_import_notice(self):
        notice = tk.Toplevel(self.win)
        notice.title("Imported")
        notice.resizable(False, False)
        notice.geometry("340x120")
        tk.Label(notice,
                 text=f"Imported from: {self.import_source}",
                 font=("Courier New", 9, "bold")).pack(pady=(20, 4))
        tk.Label(notice,
                 text="Save as a .fwd project to continue.",
                 font=("Courier New", 9)).pack()
        tk.Button(notice, text="OK", font=("Courier New", 9),
                  command=notice.destroy, width=8,
                  relief=tk.RAISED).pack(pady=12)
        notice.transient(self.win)
        notice.grab_set()

    def _build_menu(self):
        t = get_theme(self.settings)
        self.menubar = tk.Menu(self.win, bg=t["menu_bg"], fg=t["menu_fg"],
                               font=("Courier New", 9))

        file_menu = tk.Menu(self.menubar, tearoff=0,
                            bg=t["menu_bg"], fg=t["menu_fg"],
                            font=("Courier New", 9))
        file_menu.add_command(label="New Project", command=self._menu_new)
        file_menu.add_command(label="Open Project", command=self._menu_open)
        file_menu.add_command(label="Save", command=self._save,
                              accelerator="Ctrl+S")
        file_menu.add_command(label="Save As...", command=self._save_as)
        file_menu.add_separator()
        export_menu = tk.Menu(file_menu, tearoff=0,
                              bg=t["menu_bg"], fg=t["menu_fg"],
                              font=("Courier New", 9))
        export_menu.add_command(label="TXT", command=self._export_txt)
        export_menu.add_command(label="DOCX", command=self._export_docx)
        file_menu.add_cascade(label="Export", menu=export_menu)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self._on_close)
        self.menubar.add_cascade(label="File", menu=file_menu)

        settings_menu = tk.Menu(self.menubar, tearoff=0,
                                bg=t["menu_bg"], fg=t["menu_fg"],
                                font=("Courier New", 9))
        settings_menu.add_command(label="Window Size...",
                                  command=self._set_window_size)
        theme_menu = tk.Menu(settings_menu, tearoff=0,
                             bg=t["menu_bg"], fg=t["menu_fg"],
                             font=("Courier New", 9))
        theme_menu.add_command(label="Light",
                               command=lambda: self._set_theme("light"))
        theme_menu.add_command(label="Dark — Green",
                               command=lambda: self._set_theme("dark", "green"))
        theme_menu.add_command(label="Dark — Amber",
                               command=lambda: self._set_theme("dark", "amber"))
        settings_menu.add_cascade(label="Theme", menu=theme_menu)
        self.menubar.add_cascade(label="Settings", menu=settings_menu)

        self.win.config(menu=self.menubar)
        self.win.bind("<Control-s>", lambda e: self._save())

    def _build_toolbar(self):
        t = get_theme(self.settings)
        self.toolbar = tk.Frame(self.win, bg=t["bg"], relief=tk.RAISED, bd=1)
        self.toolbar.pack(fill=tk.X, side=tk.TOP)

        self.mode_btn = tk.Button(self.toolbar, text="",
                                  font=("Courier New", 9, "bold"),
                                  relief=tk.RAISED, bd=2,
                                  bg=t["btn_bg"], fg=t["btn_fg"],
                                  command=self._toggle_mode)
        self.mode_btn.pack(side=tk.LEFT, padx=6, pady=3)

    def _build_content(self):
        t = get_theme(self.settings)
        self.content_frame = tk.Frame(self.win, bg=t["bg"])
        self.content_frame.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)

        self.scrollbar = tk.Scrollbar(self.content_frame)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.text_area = tk.Text(self.content_frame,
                                 font=("Courier New", 16),
                                 wrap=tk.WORD, relief=tk.SUNKEN, bd=2,
                                 padx=24, pady=24,
                                 undo=True,
                                 yscrollcommand=self.scrollbar.set)
        self.text_area.pack(fill=tk.BOTH, expand=True)
        self.scrollbar.config(command=self.text_area.yview)

        self.text_area.tag_config("hidden")
        self.text_area.tag_config("visible")

    def _build_statusbar(self):
        t = get_theme(self.settings)
        self.statusbar = tk.Frame(self.win, bg=t["status_bg"],
                                  relief=tk.SUNKEN, bd=1, height=20)
        self.statusbar.pack(fill=tk.X, side=tk.BOTTOM)
        self.statusbar.pack_propagate(False)

        self.status_mode = tk.Label(self.statusbar, text="",
                                    font=("Courier New", 8),
                                    bg=t["status_bg"], fg=t["status_fg"],
                                    anchor=tk.W)
        self.status_mode.pack(side=tk.LEFT, padx=8)

        self.status_words = tk.Label(self.statusbar, text="",
                                     font=("Courier New", 8),
                                     bg=t["status_bg"], fg=t["status_fg"],
                                     anchor=tk.E)
        self.status_words.pack(side=tk.RIGHT, padx=8)

        self.status_file = tk.Label(self.statusbar, text="",
                                    font=("Courier New", 8),
                                    bg=t["status_bg"], fg=t["status_fg"],
                                    anchor=tk.CENTER)
        self.status_file.pack(side=tk.LEFT, expand=True)

    # ── Theme ─────────────────────────────────────────────────────────────────

    def _apply_theme(self):
        t = get_theme(self.settings)
        self.win.config(bg=t["bg"])
        self.toolbar.config(bg=t["bg"])
        self.mode_btn.config(bg=t["btn_bg"], fg=t["btn_fg"])
        self.content_frame.config(bg=t["bg"])
        self.text_area.config(bg=t["content_bg"], fg=t["content_fg"],
                              insertbackground=t["content_fg"],
                              selectbackground=t["content_fg"],
                              selectforeground=t["content_bg"])
        self.text_area.tag_config("hidden", foreground=t["hidden_fg"])
        self.text_area.tag_config("visible", foreground=t["content_fg"])
        self.statusbar.config(bg=t["status_bg"])
        self.status_mode.config(bg=t["status_bg"], fg=t["status_fg"])
        self.status_words.config(bg=t["status_bg"], fg=t["status_fg"])
        self.status_file.config(bg=t["status_bg"], fg=t["status_fg"])
        self.win.config(menu="")
        self._build_menu()

    def _set_theme(self, theme, color=None):
        self.settings["theme"] = theme
        if color:
            self.settings["dark_color"] = color
        self._apply_theme()
        if self._is_saved():
            save_fwd(self.fwd_path, self._current_content(), self.settings)

    # ── Modes ─────────────────────────────────────────────────────────────────

    def _enter_focus(self):
        self.mode = "focus"
        self.session_text = ""
        self.text_area.config(state=tk.DISABLED, cursor="arrow")
        self.mode_btn.config(text="Switch to Review")
        self.status_mode.config(text="FOCUS")
        self._refresh_focus()
        self.win.bind("<Key>", self._on_key_focus)
        self.win.bind("<Button-1>", lambda e: self.win.focus_force())
        self.win.focus_force()

    def _enter_review(self):
        self.mode = "review"
        self.text_area.config(state=tk.NORMAL, cursor="xterm")
        self.text_area.delete("1.0", tk.END)
        self.text_area.insert(tk.END, self.content)
        self.text_area.see(tk.END)
        self.mode_btn.config(text="Switch to Focus")
        self.status_mode.config(text="REVIEW")
        self._update_status()
        self.win.unbind("<Key>")
        self.win.unbind("<Button-1>")
        self.text_area.bind("<KeyRelease>", self._on_review_key)
        self.text_area.focus_set()

    def _toggle_mode(self):
        if self.mode == "focus":
            self.content += self.session_text
            self.session_text = ""
            if self._is_saved():
                save_fwd(self.fwd_path, self.content, self.settings)
            self._enter_review()
        else:
            self.content = self.text_area.get("1.0", tk.END).rstrip("\n")
            if self._is_saved():
                save_fwd(self.fwd_path, self.content, self.settings)
            self._enter_focus()

    # ── Focus mode rendering ──────────────────────────────────────────────────

    def _full_text(self):
        return self.content + self.session_text

    def _tokenize(self, text):
        import re
        result = []
        # split on newlines or whitespace, keeping delimiters
        for part in re.split("(\n+|[ \t]+)", text):
            if not part:
                continue
            is_word = bool(part.strip()) and ("\n" not in part)
            result.append((part, is_word))
        return result


    def _refresh_focus(self):
        full = self._full_text()
        tokens = self._tokenize(full)
        n = self.settings["window_size"]

        word_indices = [i for i, (t, is_word) in enumerate(tokens) if is_word]
        total_words = len(word_indices)

        if total_words <= n:
            split_idx = 0
        else:
            split_idx = word_indices[-n]

        hidden_text = "".join(t for t, _ in tokens[:split_idx])
        visible_text = "".join(t for t, _ in tokens[split_idx:])
        display_text = hidden_text + visible_text
        hidden_end = len(hidden_text)

        self.text_area.config(state=tk.NORMAL)
        self.text_area.delete("1.0", tk.END)
        self.text_area.insert(tk.END, display_text)

        self.text_area.tag_remove("hidden", "1.0", tk.END)
        self.text_area.tag_remove("visible", "1.0", tk.END)

        if hidden_end > 0:
            self.text_area.tag_add("hidden", "1.0", f"1.0 + {hidden_end} chars")
            self.text_area.tag_add("visible", f"1.0 + {hidden_end} chars", tk.END)
        else:
            self.text_area.tag_add("visible", "1.0", tk.END)

        self.text_area.mark_set("insert", tk.END)
        self.text_area.see(tk.END)
        self.text_area.config(state=tk.NORMAL)
        self.text_area.config(insertbackground=get_theme(self.settings)["content_fg"])
        self.text_area.config(state=tk.DISABLED)
        self._update_status()

    def _on_key_focus(self, event):
        blocked = {"BackSpace", "Delete", "Left", "Right", "Up", "Down",
                   "Home", "End", "Prior", "Next"}
        if event.keysym in blocked:
            return "break"
        if event.state & 0x4:
            if event.keysym.lower() == 'z':
                return "break"
            return
        if event.char and event.char.isprintable():
            self.session_text += event.char
            self._refresh_focus()
            return "break"
        if event.keysym == "Return":
            self.session_text += "\n"
            self._refresh_focus()
            return "break"
        if event.keysym == "Tab":
            self.session_text += "\t"
            self._refresh_focus()
            return "break"
        if event.keysym == "space":
            self.session_text += " "
            self._refresh_focus()
            return "break"
        return "break"

    def _on_review_key(self, event):
        self._update_status()

    # ── Status bar ────────────────────────────────────────────────────────────

    def _update_status(self):
        text = self._current_content()
        words = len(text.split()) if text.strip() else 0
        self.status_words.config(text=f"{words} words")
        fname = os.path.basename(self.fwd_path) if self.fwd_path else "Untitled"
        self.status_file.config(text=fname)

    # ── Save ──────────────────────────────────────────────────────────────────

    def _save(self):
        if not self._is_saved():
            self._save_as()
        else:
            save_fwd(self.fwd_path, self._current_content(), self.settings)

    def _save_as(self):
        # Suggest filename based on import source or untitled
        if self.import_source:
            suggested = os.path.splitext(self.import_source)[0] + FWD_EXT
        else:
            suggested = "Untitled" + FWD_EXT

        path = filedialog.asksaveasfilename(
            title="Save project as",
            defaultextension=FWD_EXT,
            initialfile=suggested,
            filetypes=[(f"ForwardOnly files", f"*{FWD_EXT}")])
        if not path:
            return
        self.fwd_path = path
        self.content = self._current_content()
        self.import_source = None
        save_fwd(self.fwd_path, self.content, self.settings)
        fname = os.path.basename(self.fwd_path)
        self.win.title(f"{APP_NAME} — {fname}")
        self._update_status()

    # ── Export ────────────────────────────────────────────────────────────────

    def _export_txt(self):
        content = self._current_content()
        path = export_txt(content, self.fwd_path,
                          self.settings.get("export_txt", ""))
        if path:
            self.settings["export_txt"] = path
            if self._is_saved():
                save_fwd(self.fwd_path, self._current_content(), self.settings)

    def _export_docx(self):
        content = self._current_content()
        path = export_docx(content, self.fwd_path,
                           self.settings.get("export_docx", ""))
        if path:
            self.settings["export_docx"] = path
            if self._is_saved():
                save_fwd(self.fwd_path, self._current_content(), self.settings)

    # ── Settings ──────────────────────────────────────────────────────────────

    def _set_window_size(self):
        val = simpledialog.askinteger(
            "Window Size",
            "Number of visible words in Focus mode:",
            initialvalue=self.settings["window_size"],
            minvalue=1, maxvalue=50,
            parent=self.win)
        if val:
            self.settings["window_size"] = val
            if self._is_saved():
                save_fwd(self.fwd_path, self._current_content(), self.settings)
            if self.mode == "focus":
                self._refresh_focus()

    # ── File menu ─────────────────────────────────────────────────────────────

    def _menu_new(self):
        if not self._prompt_save_before_leaving():
            return
        self.win.destroy()
        self.fwd_path = None
        self.content = ""
        self.settings = dict(DEFAULT_SETTINGS)
        self.session_text = ""
        self.import_source = None
        self._show_launcher()

    def _menu_open(self):
        if not self._prompt_save_before_leaving():
            return
        self.win.destroy()
        self._show_launcher()

    # ── Close ─────────────────────────────────────────────────────────────────

    def _on_close(self):
        if not self._is_saved() and self._has_content():
            answer = messagebox.askyesnocancel(
                "Save before closing?",
                "This project hasn't been saved yet.\nSave now?")
            if answer is None:
                return
            if answer:
                self._save_as()
                if not self._is_saved():
                    return
        elif self._is_saved():
            if self.mode == "focus" and self.session_text:
                self._save()
            elif self.mode == "review":
                self._save()
        self.root.destroy()


def main():
    import sys
    open_path = None
    if len(sys.argv) > 1:
        arg = sys.argv[1]
        if os.path.isfile(arg) and arg.lower().endswith(FWD_EXT):
            open_path = arg

    root = tk.Tk()
    root.withdraw()
    app = ForwardOnly(root, open_path=open_path)
    root.mainloop()

if __name__ == "__main__":
    main()
